import csv
import json
import zipfile
import io
from dataclasses import dataclass, asdict
from pathlib import Path
from typing import Any, List, Optional
import re

import fitz
import pdfplumber
from docx import Document
from openpyxl import load_workbook
from PIL import Image, ImageOps, ImageFilter
import pytesseract
from pytesseract import Output
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".xlsx",
    ".csv",
    ".txt",
    ".png",
    ".jpg",
    ".jpeg",
}


@dataclass
class ExtractedTable:
    source_file: str
    source_type: str
    page_or_sheet: str
    table_index: int
    rows: List[List[str]]


@dataclass
class ExtractedDocument:
    source_file: str
    source_type: str
    text: str
    tables: List[ExtractedTable]
    warnings: List[str]


def clean_cell(value: Any) -> str:
    """
    Converts a table cell value into clean text.

    Why this exists:
    - PDF/Excel/Word cells can contain None, numbers, dates, or messy spacing.
    - Gemini needs readable evidence, so every cell is converted into a clean string.
    """
    if value is None:
        return ""

    text = str(value)
    text = text.replace("\n", " ")
    text = " ".join(text.split())
    return text.strip()


def table_to_markdown(table: ExtractedTable) -> str:
    """
    Converts extracted table rows into markdown.

    Why markdown:
    - It is readable in logs and Word validation files.
    - Gemini understands markdown tables better than raw Python lists.
    """
    rows = table.rows

    if not rows:
        return ""

    max_columns = max(len(row) for row in rows)
    normalized_rows = []

    for row in rows:
        padded_row = row + [""] * (max_columns - len(row))
        normalized_rows.append(padded_row)

    header = normalized_rows[0]
    separator = ["---"] * max_columns
    body = normalized_rows[1:]

    lines = []
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join(separator) + " |")

    for row in body:
        lines.append("| " + " | ".join(row) + " |")

    return "\n".join(lines)

def normalize_ocr_table_decimals(rows: List[List[str]]) -> List[List[str]]:
    """
    Fixes OCR decimal loss for known decimal-style columns.

    Example:
    - FG% column: 515 -> .515
    - AVG column: 282 -> 28.2

    Important:
    This runs after the header is detected, so it does not damage years,
    games, rebounds, points, or other whole-number columns.
    """
    if not rows:
        return rows

    header = [cell.strip().lower().replace(".", "") for cell in rows[0]]

    decimal_columns = {}

    for index, column_name in enumerate(header):
        if column_name in {"fg%", "fg", "pct", "percentage"}:
            decimal_columns[index] = "leading_decimal"

        if column_name in {"avg", "average"}:
            decimal_columns[index] = "one_decimal"

    normalized_rows = [rows[0]]

    for row in rows[1:]:
        new_row = row[:]

        for index, decimal_type in decimal_columns.items():
            if index >= len(new_row):
                continue

            value = new_row[index].strip()

            if decimal_type == "leading_decimal":
                # 515 -> .515, 482 -> .482
                if re.fullmatch(r"\d{3}", value):
                    new_row[index] = f".{value}"

            elif decimal_type == "one_decimal":
                # 282 -> 28.2, 350 -> 35.0, 227 -> 22.7
                if re.fullmatch(r"\d{3}", value):
                    new_row[index] = f"{value[:-1]}.{value[-1]}"

        normalized_rows.append(new_row)

    return normalized_rows

def is_noise_ocr_row(row: List[str]) -> bool:
    """
    Removes OCR noise rows that are not real document content.

    Example removed:
    - (x=15, y=38) ~ R:255 G:254 B:251
    """
    joined = " ".join(row).lower()

    if "x=" in joined and "y=" in joined:
        return True

    if "r:255" in joined or "g:254" in joined or "b:251" in joined:
        return True

    noise_patterns = [
        "r:",
        "g:",
        "b:",
        "x=",
        "y=",
    ]

    return any(pattern in joined for pattern in noise_patterns)

def preprocess_ocr_image(image: Image.Image) -> Image.Image:
    """
    Prepares an image for OCR without destroying small table text.

    Why:
    - Hard black/white thresholding can erase thin numbers or grid text.
    - For scanned tables, grayscale + enlargement + sharpening is safer.
    """
    grayscale = ImageOps.grayscale(image)

    enlarged = grayscale.resize(
        (grayscale.width * 2, grayscale.height * 2)
    )

    sharpened = enlarged.filter(ImageFilter.SHARPEN)

    return sharpened


def ocr_image_to_text_and_rows(image: Image.Image) -> tuple[str, List[List[str]]]:
    """
    Runs OCR and also tries to reconstruct table-like rows.

    Normal OCR text:
    - Good for paragraphs.

    OCR row reconstruction:
    - Uses Tesseract word bounding boxes.
    - Groups words by detected line number.
    - This is better for scanned tables because it preserves row-level structure.
    """
    processed_image = preprocess_ocr_image(image)

    text_config = "--psm 6 -c preserve_interword_spaces=1"
    ocr_text = clean_ocr_text(
        pytesseract.image_to_string(processed_image, config=text_config).strip()
    )   

    data = pytesseract.image_to_data(
        processed_image,
        output_type=Output.DICT,
        config="--psm 6",
    )

    line_groups: dict[tuple[int, int, int], List[tuple[int, str]]] = {}

    for index, raw_text in enumerate(data.get("text", [])):
        word = clean_ocr_token(clean_cell(raw_text))

        if not word:
            continue

        try:
            confidence = float(data["conf"][index])
        except ValueError:
            confidence = -1

        if confidence < 25:
            continue

        block_num = data["block_num"][index]
        paragraph_num = data["par_num"][index]
        line_num = data["line_num"][index]
        left = data["left"][index]

        key = (block_num, paragraph_num, line_num)

        if key not in line_groups:
            line_groups[key] = []

        line_groups[key].append((left, word))

    rows: List[List[str]] = []

    for key in sorted(line_groups.keys()):
        words = sorted(line_groups[key], key=lambda item: item[0])

        original_row = [word for _, word in words]

        # Check noise before cleaning, so rows like
        # (x=15, y=38) ~ R:255 G:254 B:251 are removed properly.
        if is_noise_ocr_row(original_row):
            continue

        cleaned_row = [
            clean_ocr_token(word)
            for word in original_row
        ]

        cleaned_row = [
            word
            for word in cleaned_row
            if word
        ]

        if cleaned_row and not is_noise_ocr_row(cleaned_row):
            rows.append(cleaned_row)

    return ocr_text, rows

def clean_ocr_table_rows(rows: List[List[str]]) -> List[List[str]]:
    """
    Cleans OCR table rows before saving them.

    Why:
    - OCR may include title rows like 'Table'.
    - OCR may include pixel/color metadata from screenshots.
    - OCR may include blank cells or rows before the real header.
    - We want the saved OCR table to begin from the real table header.
    """
    cleaned_rows: List[List[str]] = []

    for row in rows:
        if is_noise_ocr_row(row):
            continue

        non_empty_cells = [
            cell.strip().strip("|")
            for cell in row
            if cell.strip().strip("|")
        ]

        if not non_empty_cells:
            continue

        joined = " ".join(non_empty_cells).lower()

        # Remove simple title-only rows like "Table" or "Table |"
        if joined in {"table", ":", ";"}:
            continue

        cleaned_rows.append(non_empty_cells)

    # If we can identify a real header row, discard everything before it.
    table_header_keywords = {
        "year",
        "club",
        "control",
        "requirement",
        "result",
        "owner",
        "status",
        "finding",
        "risk",
        "amount",
        "balance",
        "date",
        "pts",
        "avg",
        "reb",
        "ast",
    }

    header_start_index = None

    for index, row in enumerate(cleaned_rows):
        row_words = {cell.lower().replace(".", "") for cell in row}
        matching_keywords = row_words.intersection(table_header_keywords)
        

        if len(matching_keywords) >= 2:
            header_start_index = index
            break

    if header_start_index is not None:
        cleaned_rows = cleaned_rows[header_start_index:]

    return cleaned_rows


def is_number_like(value: str) -> bool:
    """
    Checks whether a token looks numeric.

    Examples:
    - 1984-85
    - 82
    - 51.5
    - 19%
    - 14016
    """
    cleaned = value.replace(",", "").replace(".", "").replace("%", "").replace("-", "")
    return cleaned.isdigit()

def clean_ocr_text(text: str) -> str:
    """
    Removes obvious OCR/debug noise from raw OCR text.
    """
    cleaned_lines = []

    for line in text.splitlines():
        lower_line = line.lower()

        if "r:255" in lower_line or "g:254" in lower_line or "b:251" in lower_line:
            continue

        if "x=" in lower_line and "y=" in lower_line:
            continue

        cleaned_lines.append(line)

    return "\n".join(cleaned_lines).strip()

def looks_like_ocr_table(rows: List[List[str]]) -> bool:
    """
    Decides whether OCR rows look like a real table.

    Why stricter:
    - Paragraph text also becomes OCR rows.
    - We only want to save an OCR table if it has table-like headers
      or repeated numeric/structured rows.
    """
    if not rows:
        return False

    flattened_text = " ".join(" ".join(row) for row in rows).lower()

    table_keywords = [
        "year",
        "club",
        "pts",
        "avg",
        "reb",
        "ast",
        "stl",
        "blk",
        "control",
        "result",
        "owner",
        "status",
        "finding",
        "risk",
        "amount",
        "balance",
        "date",
    ]

    has_table_header = any(keyword in flattened_text for keyword in table_keywords)

    numeric_heavy_rows = 0

    for row in rows:
        if len(row) < 4:
            continue

        numeric_count = sum(1 for cell in row if is_number_like(cell))

        if numeric_count >= 2:
            numeric_heavy_rows += 1

    return has_table_header and numeric_heavy_rows >= 1

def ocr_pdf_page(
    pdf_document: fitz.Document,
    page_index: int,
    dpi: int = 360,
) -> tuple[str, List[List[str]]]:
    """
    Converts one PDF page into an image and runs OCR on it.

    Returns:
    - OCR text for paragraphs.
    - OCR rows for scanned/image-only tables.
    """
    zoom = dpi / 72
    matrix = fitz.Matrix(zoom, zoom)

    page = pdf_document.load_page(page_index)
    pixmap = page.get_pixmap(matrix=matrix, alpha=False)

    image_bytes = pixmap.tobytes("png")
    image = Image.open(io.BytesIO(image_bytes))

    return ocr_image_to_text_and_rows(image)

def clean_ocr_token(token: str) -> str:
    """
    Cleans OCR garbage from individual OCR tokens without destroying valid values.

    Examples:
    - '~=—s«65' becomes '65'
    - '#£=x387' becomes '387'
    - '1984-85' stays '1984-85'
    - '.457' stays '.457'
    - '37.1' stays '37.1'
    - '4°27' becomes '427'
    """
    token = clean_cell(token)

    if not token:
        return ""

    # Preserve year/season ranges like 1984-85.
    if re.fullmatch(r"\d{4}-\d{2}", token):
        return token

    # Preserve normal numbers, decimals, leading-dot decimals, and percentages.
    if re.fullmatch(r"-?(?:\d+|\d*\.\d+)%?", token):
        return token

    # Fix OCR degree symbol inside numbers, for example 4°27 -> 427.
    token_without_degree = re.sub(r"(?<=\d)[°º](?=\d)", "", token)

    if re.fullmatch(r"-?(?:\d+|\d*\.\d+)%?", token_without_degree):
        return token_without_degree

    # If the token is mostly OCR junk but contains a number, keep the last useful number.
    # Example: '~=—s«65' -> '65', '#£=x387' -> '387'
    number_matches = re.findall(r"-?(?:\d*\.\d+|\d+)%?", token_without_degree)

    if number_matches and re.search(r"[^A-Za-z0-9.\-%]", token_without_degree):
        return number_matches[-1]

    # Otherwise remove only weird symbols, not normal table text.
    cleaned = re.sub(r"[^A-Za-z0-9.,:%()/\-]", "", token_without_degree)

    return cleaned.strip()

def extract_pdf(file_path: Path) -> ExtractedDocument:
    """
    Extracts text and tables from a PDF.

    Logic:
    1. First try selectable text extraction using pdfplumber.
    2. If a page has no selectable text, run OCR fallback on that page.
    3. If OCR rows look table-like, save them as an OCR-derived table.
    4. Still try pdfplumber table extraction for machine-readable PDF tables.

    Why this design:
    - Direct PDF extraction is cleaner and faster for normal PDFs.
    - OCR is used only when the PDF is scanned/image-only.
    - OCR table reconstruction is separate because plain OCR does not preserve table structure well.
    """
    text_parts: List[str] = []
    tables: List[ExtractedTable] = []
    warnings: List[str] = []

    try:
        with pdfplumber.open(file_path) as pdf, fitz.open(file_path) as ocr_pdf:
            for page_number, page in enumerate(pdf.pages, start=1):
                page_text = page.extract_text() or ""

                if page_text.strip():
                    text_parts.append(
                        f"\n\n--- Page {page_number} Selectable Text ---\n{page_text}"
                    )
                else:
                    warnings.append(
                        f"Page {page_number}: No selectable text found. OCR fallback used."
                    )

                    try:
                        ocr_text, ocr_rows = ocr_pdf_page(
                            pdf_document=ocr_pdf,
                            page_index=page_number - 1,
                        )

                        if ocr_text:
                            text_parts.append(
                                f"\n\n--- Page {page_number} OCR Text ---\n{ocr_text}"
                            )
                        else:
                            warnings.append(
                                f"Page {page_number}: OCR ran but returned no text."
                            )

                        cleaned_ocr_rows = clean_ocr_table_rows(ocr_rows)

                        if looks_like_ocr_table(ocr_rows):
                            table_rows = clean_ocr_table_rows(ocr_rows)
                            table_rows = remove_ocr_noise_rows(table_rows)
                            table_rows = normalize_ocr_table_decimals(table_rows)
                            table_rows = remove_ocr_noise_rows(table_rows)

                            if table_rows:
                                tables.append(
                                    ExtractedTable(
                                        source_file=file_path.name,
                                        source_type="pdf",
                                        page_or_sheet=f"page {page_number} OCR",
                                        table_index=len(tables) + 1,
                                        rows=table_rows,
                                    )
                                )

                    except Exception as ocr_error:
                        warnings.append(
                            f"Page {page_number}: OCR fallback failed: {ocr_error}"
                        )

                try:
                    page_tables = page.extract_tables() or []

                    for table_index, raw_table in enumerate(page_tables, start=1):
                        cleaned_rows = [
                            [clean_cell(cell) for cell in row]
                            for row in raw_table
                            if row is not None
                        ]

                        if cleaned_rows:
                            tables.append(
                                ExtractedTable(
                                    source_file=file_path.name,
                                    source_type="pdf",
                                    page_or_sheet=f"page {page_number}",
                                    table_index=table_index,
                                    rows=cleaned_rows,
                                )
                            )

                except Exception as table_error:
                    warnings.append(
                        f"Page {page_number}: Table extraction failed: {table_error}"
                    )

    except Exception as error:
        warnings.append(f"PDF extraction failed: {error}")

    return ExtractedDocument(
        source_file=file_path.name,
        source_type="pdf",
        text="\n".join(text_parts).strip(),
        tables=tables,
        warnings=warnings,
    )

def remove_ocr_noise_rows(rows: List[List[str]]) -> List[List[str]]:
    """
    Removes OCR noise rows that accidentally survive cleaning/normalization.

    Example removed:
    - 15 | 38 | 255 | 254 | 251
    - 15 | 38 | 255 | .254 | 251
    """
    if not rows:
        return rows

    cleaned_rows = [rows[0]]

    for row in rows[1:]:
        compact = [cell.strip() for cell in row if cell.strip()]

        # First use the existing noise detector.
        if is_noise_ocr_row(compact):
            continue

        normalized = []

        for cell in compact:
            value = cell.strip()

            # Convert .254 back to 254 only for noise detection.
            if value.startswith(".") and value[1:].isdigit():
                value = value[1:]

            value = value.replace(".", "")

            normalized.append(value)

        # Catch coordinate/RGB rows after OCR cleanup.
        if len(normalized) == 5 and all(value.isdigit() for value in normalized):
            numbers = [int(value) for value in normalized]

            looks_like_coordinate_rgb_row = (
                0 <= numbers[0] <= 5000
                and 0 <= numbers[1] <= 5000
                and 0 <= numbers[2] <= 255
                and 0 <= numbers[3] <= 255
                and 0 <= numbers[4] <= 255
            )

            if looks_like_coordinate_rgb_row:
                continue

        cleaned_rows.append(row)

    return cleaned_rows

def extract_docx(file_path: Path) -> ExtractedDocument:
    """
    Extracts paragraphs and real Word tables from DOCX files.

    Why python-docx:
    - It reads the actual Word document structure.
    - Tables remain as rows and cells instead of messy flat text.
    """
    text_parts: List[str] = []
    tables: List[ExtractedTable] = []
    warnings: List[str] = []

    try:
        document = Document(file_path)

        paragraph_text = [
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]

        if paragraph_text:
            text_parts.append("\n".join(paragraph_text))

        for table_index, table in enumerate(document.tables, start=1):
            rows: List[List[str]] = []

            for row in table.rows:
                rows.append([clean_cell(cell.text) for cell in row.cells])

            if rows:
                tables.append(
                    ExtractedTable(
                        source_file=file_path.name,
                        source_type="docx",
                        page_or_sheet="document",
                        table_index=table_index,
                        rows=rows,
                    )
                )

    except Exception as error:
        warnings.append(f"DOCX extraction failed: {error}")

    return ExtractedDocument(
        source_file=file_path.name,
        source_type="docx",
        text="\n".join(text_parts).strip(),
        tables=tables,
        warnings=warnings,
    )


def extract_xlsx(file_path: Path) -> ExtractedDocument:
    """
    Extracts workbook sheets as structured tables.

    Why openpyxl:
    - It reads Excel files directly.
    - It preserves workbook -> sheet -> row -> cell structure.
    - This is easier to explain in code review than hidden dataframe logic.
    """
    text_parts: List[str] = []
    tables: List[ExtractedTable] = []
    warnings: List[str] = []

    try:
        workbook = load_workbook(file_path, data_only=True)

        for sheet in workbook.worksheets:
            rows: List[List[str]] = []

            for row in sheet.iter_rows(values_only=True):
                cleaned_row = [clean_cell(cell) for cell in row]

                if any(cleaned_row):
                    rows.append(cleaned_row)

            if rows:
                tables.append(
                    ExtractedTable(
                        source_file=file_path.name,
                        source_type="xlsx",
                        page_or_sheet=sheet.title,
                        table_index=1,
                        rows=rows,
                    )
                )

                preview_lines = [
                    "\t".join(row)
                    for row in rows[:20]
                ]

                text_parts.append(
                    f"\n--- Sheet: {sheet.title} Preview ---\n"
                    + "\n".join(preview_lines)
                )

    except Exception as error:
        warnings.append(f"XLSX extraction failed: {error}")

    return ExtractedDocument(
        source_file=file_path.name,
        source_type="xlsx",
        text="\n".join(text_parts).strip(),
        tables=tables,
        warnings=warnings,
    )


def extract_csv(file_path: Path) -> ExtractedDocument:
    """
    Extracts CSV rows as one structured table.

    Why simple csv module:
    - CSV files are already table data.
    - The built-in csv module is enough and easy to explain.
    """
    tables: List[ExtractedTable] = []
    warnings: List[str] = []

    try:
        with file_path.open("r", encoding="utf-8-sig", newline="") as file:
            reader = csv.reader(file)
            rows = [[clean_cell(cell) for cell in row] for row in reader]

        rows = [row for row in rows if any(row)]

        if rows:
            tables.append(
                ExtractedTable(
                    source_file=file_path.name,
                    source_type="csv",
                    page_or_sheet="csv",
                    table_index=1,
                    rows=rows,
                )
            )

        preview = "\n".join([", ".join(row) for row in rows[:30]])

    except UnicodeDecodeError:
        try:
            with file_path.open("r", encoding="latin-1", newline="") as file:
                reader = csv.reader(file)
                rows = [[clean_cell(cell) for cell in row] for row in reader]

            rows = [row for row in rows if any(row)]

            if rows:
                tables.append(
                    ExtractedTable(
                        source_file=file_path.name,
                        source_type="csv",
                        page_or_sheet="csv",
                        table_index=1,
                        rows=rows,
                    )
                )

            preview = "\n".join([", ".join(row) for row in rows[:30]])

        except Exception as error:
            warnings.append(f"CSV extraction failed: {error}")
            preview = ""

    except Exception as error:
        warnings.append(f"CSV extraction failed: {error}")
        preview = ""

    return ExtractedDocument(
        source_file=file_path.name,
        source_type="csv",
        text=preview.strip(),
        tables=tables,
        warnings=warnings,
    )


def extract_txt(file_path: Path) -> ExtractedDocument:
    """
    Extracts plain text files.

    Why this exists:
    - README files, notes, and manifest files are often TXT.
    - These are already readable and do not need complex parsing.
    """
    warnings: List[str] = []

    try:
        text = file_path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        text = file_path.read_text(encoding="latin-1")
    except Exception as error:
        warnings.append(f"TXT extraction failed: {error}")
        text = ""

    return ExtractedDocument(
        source_file=file_path.name,
        source_type="txt",
        text=text.strip(),
        tables=[],
        warnings=warnings,
    )


def extract_image_ocr(file_path: Path) -> ExtractedDocument:
    """
    Extracts text from image files using OCR.

    Why pytesseract:
    - It runs locally.
    - It is free/open-source.
    - Good enough for clean screenshots or scanned text pages.

    Important:
    - On Windows, pytesseract needs the Tesseract desktop app installed.
    - If OCR fails, check whether tesseract.exe is installed and in PATH.
    """
    warnings: List[str] = []

    try:
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image)
    except Exception as error:
        warnings.append(f"OCR extraction failed: {error}")
        text = ""

    return ExtractedDocument(
        source_file=file_path.name,
        source_type="image_ocr",
        text=text.strip(),
        tables=[],
        warnings=warnings,
    )


def extract_single_file(file_path: str | Path) -> ExtractedDocument:
    """
    Routes one file to the correct extraction function.

    This is the main decision function:
    - PDF -> pdfplumber
    - DOCX -> python-docx
    - XLSX -> openpyxl
    - CSV -> csv module
    - TXT -> normal text read
    - image -> OCR
    """
    path = Path(file_path)
    extension = path.suffix.lower()

    if extension == ".pdf":
        return extract_pdf(path)

    if extension == ".docx":
        return extract_docx(path)

    if extension == ".xlsx":
        return extract_xlsx(path)

    if extension == ".csv":
        return extract_csv(path)

    if extension == ".txt":
        return extract_txt(path)

    if extension in {".png", ".jpg", ".jpeg"}:
        return extract_image_ocr(path)

    return ExtractedDocument(
        source_file=path.name,
        source_type="unsupported",
        text="",
        tables=[],
        warnings=[f"Unsupported file type: {extension}"],
    )


def extract_zip_package(zip_path: str | Path, output_folder: str | Path) -> List[ExtractedDocument]:
    """
    Extracts all supported files from a ZIP package.

    Why this exists:
    - Your app accepts audit ZIP packages.
    - The ZIP may contain PDFs, Word reports, Excel files, CSVs, TXT files, and images.
    """
    zip_path = Path(zip_path)
    output_folder = Path(output_folder)
    output_folder.mkdir(parents=True, exist_ok=True)

    extracted_documents: List[ExtractedDocument] = []

    with zipfile.ZipFile(zip_path, "r") as zip_file:
        zip_file.extractall(output_folder)

    for file_path in output_folder.rglob("*"):
        if file_path.is_file() and file_path.suffix.lower() in SUPPORTED_EXTENSIONS:
            extracted_documents.append(extract_single_file(file_path))

    return extracted_documents


def build_readable_report(extracted_documents: List[ExtractedDocument]) -> str:
    """
    Converts extracted documents into one readable report.

    This report is useful for:
    - Monday validation
    - debugging extraction quality
    - checking what Gemini will receive
    """
    report_parts: List[str] = []

    for document in extracted_documents:
        report_parts.append("=" * 90)
        report_parts.append(f"FILE: {document.source_file}")
        report_parts.append(f"TYPE: {document.source_type}")
        report_parts.append("=" * 90)

        if document.warnings:
            report_parts.append("\nWARNINGS:")
            for warning in document.warnings:
                report_parts.append(f"- {warning}")

        if document.text:
            report_parts.append("\nEXTRACTED TEXT:")
            report_parts.append(document.text[:8000])

        if document.tables:
            report_parts.append("\nEXTRACTED TABLES:")

            for table in document.tables:
                report_parts.append(
                    f"\nTable {table.table_index} from {table.page_or_sheet}"
                )
                report_parts.append(table_to_markdown(table))

        report_parts.append("\n\n")

    return "\n".join(report_parts)


def save_extraction_outputs(
    extracted_documents: List[ExtractedDocument],
    output_folder: str | Path,
) -> None:
    """
    Saves extraction output as TXT and JSON.

    TXT:
    - Easy to paste into Word validation document.

    JSON:
    - Useful later if backend needs structured data for Gemini.
    """
    output_folder = Path(output_folder)
    output_folder.mkdir(parents=True, exist_ok=True)

    readable_report = build_readable_report(extracted_documents)

    txt_path = output_folder / "extracted_output_readable.txt"
    json_path = output_folder / "extracted_output_structured.json"

    txt_path.write_text(readable_report, encoding="utf-8")

    json_ready = []
    for document in extracted_documents:
        document_dict = asdict(document)
        json_ready.append(document_dict)

    json_path.write_text(
        json.dumps(json_ready, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )


if __name__ == "__main__":
    """
    Manual test command examples:

    Single file:
    python extraction.py "C:\\path\\to\\file.pdf" "C:\\path\\to\\validation_outputs"

    ZIP package:
    python extraction.py "C:\\path\\to\\audit_package.zip" "C:\\path\\to\\validation_outputs"
    """
    import sys

    if len(sys.argv) < 3:
        print("Usage:")
        print("python extraction.py <input_file_or_zip> <output_folder>")
        sys.exit(1)

    input_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])

    if input_path.suffix.lower() == ".zip":
        temp_extract_folder = output_path / "unzipped_files"
        documents = extract_zip_package(input_path, temp_extract_folder)
    else:
        documents = [extract_single_file(input_path)]

    save_extraction_outputs(documents, output_path)

    print(f"Extraction complete. Outputs saved to: {output_path}")